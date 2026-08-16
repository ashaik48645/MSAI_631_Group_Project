"""
MSAI 631 Course Support RAG Chatbot
Contributor: Abdul Kareem Shaik

Loads approved course documents from PDF, DOCX, and TXT files.
"""

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from config import SUPPORTED_EXTENSIONS


@dataclass
class LoadedDocument:
    source: str
    text: str


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Include text from tables because course handouts often contain them.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def load_document(file_path: str | Path) -> LoadedDocument:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            f"Supported types: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    if suffix == ".txt":
        text = _read_txt(path)
    elif suffix == ".pdf":
        text = _read_pdf(path)
    else:
        text = _read_docx(path)

    return LoadedDocument(source=path.name, text=text)


def load_documents(file_paths: Iterable[str | Path]) -> List[LoadedDocument]:
    loaded = []
    for file_path in file_paths:
        doc = load_document(file_path)
        if doc.text.strip():
            loaded.append(doc)
    return loaded
